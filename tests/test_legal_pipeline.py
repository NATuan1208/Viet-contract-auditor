from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from pipeline.connectors import (
    parse_rss_items,
    parse_thuvienphapluat_new_documents,
    parse_vanban_chinhphu_detail_record,
    parse_vanban_chinhphu_listing,
)
from pipeline.iceberg_lakehouse import _replace_netloc, bronze_row, gold_row, redact_uri, silver_row
from pipeline.lakehouse import LocalLakehouse
from pipeline.models import KGUpdateManifest, LegalDocumentRecord, RawArtifact, SourceConfig, SourceItem
from pipeline.normalization import checksum_bytes, split_articles, stable_doc_id
from pipeline.registry import load_source_registry
from pipeline.versioning import DocumentVersionStore
from pipeline.kg_update import _document_text_with_provenance, _insert_kwargs, latest_manifest_items, load_manifest
from pipeline.attachment_extraction import extract_attachment_text, fetch_attachment, merge_attachment_extractions
from core.rerank_client import (
    RerankSettings,
    _sigmoid,
    _truncate_for_rerank,
    build_rerank_model_func,
    get_rerank_settings,
)
from core.benchmark_trace import score_summary, stable_text_hash, text_preview, write_trace_event
from core.lightrag_client import query_hybrid
from benchmark_rerank_ab import (
    BenchmarkVariant,
    SWEEP_VARIANTS,
    build_summary_markdown,
    build_variant_command,
    parse_eval_metrics,
    parse_state_summary,
    parse_trace_summary,
)
from benchmark_diagnose import build_diagnosis, build_diagnosis_markdown, diagnose_result
from benchmark_calibrate import build_calibration, build_calibration_markdown


class LegalPipelineTests(unittest.TestCase):
    def test_source_registry_loads_and_blocks_commercial_detail(self) -> None:
        _, sources = load_source_registry(ROOT / "config" / "legal_sources.yml")
        self.assertGreaterEqual(len(sources), 3)
        commercial = [source for source in sources if source.discovery_only]
        self.assertTrue(commercial)
        self.assertTrue(all("detail" not in source.crawl_methods for source in commercial))

    def test_rss_parser_filters_since(self) -> None:
        source = SourceConfig(
            source_id="congbao",
            name="Cong bao",
            tier=0,
            role="canonical_update_feed",
            base_url="https://congbao.chinhphu.vn",
            domain_whitelist=["congbao.chinhphu.vn"],
            crawl_methods=["rss"],
            rate_limit_per_minute=30,
            robots_policy="obey",
            priority=10,
            license_note="attribute",
        )
        rss = """<?xml version="1.0"?>
        <rss><channel>
          <item>
            <title>Van ban moi</title>
            <link>https://congbao.chinhphu.vn/item-1</link>
            <pubDate>Wed, 13 May 2026 08:00:00 GMT</pubDate>
          </item>
          <item>
            <title>Cu</title>
            <link>https://congbao.chinhphu.vn/item-0</link>
            <pubDate>Mon, 11 May 2026 08:00:00 GMT</pubDate>
          </item>
        </channel></rss>"""
        items = parse_rss_items(
            rss,
            source,
            "test_feed",
            since=__import__("datetime").datetime.fromisoformat("2026-05-12T00:00:00+00:00"),
        )
        self.assertEqual(len(items), 1)
        self.assertEqual(items[0].title, "Van ban moi")

    def test_thuvienphapluat_parser_is_discovery_only_metadata(self) -> None:
        source = SourceConfig(
            source_id="thuvienphapluat_discovery",
            name="Thu vien Phap luat discovery signal",
            tier=2,
            role="commercial_discovery_only",
            base_url="https://thuvienphapluat.vn",
            domain_whitelist=["thuvienphapluat.vn"],
            crawl_methods=["search"],
            rate_limit_per_minute=5,
            robots_policy="obey",
            priority=1,
            license_note="Discovery only",
        )
        html = """
        <article>
          <a href="/van-ban/Bo-may-hanh-chinh/Nghi-dinh-148-2026-ND-CP-705175.aspx">
            Nghi dinh 148/2026/ND-CP cua Chinh phu ve quan ly thong tin doi ngoai
          </a>
          <span>Ban hanh: 12/05/2026</span>
        </article>
        """
        items = parse_thuvienphapluat_new_documents(html, source)
        self.assertEqual(len(items), 1)
        self.assertEqual(items[0].metadata["source_use"], "discovery_only")
        self.assertEqual(items[0].metadata["canonical_number"], "148/2026/ND-CP")
        html_short_number = """
        <article>
          <a href="/van-ban/Cong-nghe-thong-tin/Quyet-dinh-840-QD-TTg-2026-705265.aspx">
            Quyet dinh 840/QD-TTg nam 2026 phe duyet chuong trinh phat trien
          </a>
        </article>
        """
        short_items = parse_thuvienphapluat_new_documents(html_short_number, source)
        self.assertEqual(short_items[0].metadata["canonical_number"], "840/QD-TTg")

    def test_vanban_chinhphu_listing_discovers_real_document_items(self) -> None:
        source = SourceConfig(
            source_id="vanban_chinhphu",
            name="He thong van ban Chinh phu",
            tier=0,
            role="official_fulltext",
            base_url="https://vanban.chinhphu.vn",
            domain_whitelist=["vanban.chinhphu.vn"],
            crawl_methods=["search", "detail"],
            rate_limit_per_minute=20,
            robots_policy="obey",
            priority=8,
            license_note="Official source",
        )
        html = """
        <tr>
          <td><a href="/?docid=218020&pageid=27160">144/2026/ND-CP</a></td>
          <td>05/05/2026</td>
          <td><a href="/?docid=218020&pageid=27160">Sua doi mot so dieu ve thue gia tri gia tang</a></td>
        </tr>
        """
        items = parse_vanban_chinhphu_listing(html, source)
        self.assertEqual(len(items), 1)
        self.assertEqual(items[0].external_id, "218020")
        self.assertEqual(items[0].metadata["canonical_number"], "144/2026/ND-CP")
        self.assertEqual(items[0].metadata["issue_date"], "2026-05-05")

    def test_vanban_chinhphu_detail_parse_creates_legal_document_record(self) -> None:
        source = SourceConfig(
            source_id="vanban_chinhphu",
            name="He thong van ban Chinh phu",
            tier=0,
            role="official_fulltext",
            base_url="https://vanban.chinhphu.vn",
            domain_whitelist=["vanban.chinhphu.vn"],
            crawl_methods=["search", "detail"],
            rate_limit_per_minute=20,
            robots_policy="obey",
            priority=8,
            license_note="Official source",
        )
        item = SourceItem(
            source_id="vanban_chinhphu",
            url="https://vanban.chinhphu.vn/?docid=218020&pageid=27160",
            title="Nghi dinh 144/2026/ND-CP: Sua doi mot so dieu ve thue gia tri gia tang",
            metadata={"crawl_method": "official_listing"},
        )
        raw = RawArtifact(
            source_item=item,
            content=b"""
            <html><body>
              <h1>Nghi dinh so 144/2026/ND-CP cua Chinh phu</h1>
              <div>So ky hieu 144/2026/ND-CP</div>
              <div>Ngay ban hanh 05-05-2026</div>
              <div>Ngay co hieu luc 20-06-2026</div>
              <div>Loai van ban Nghi dinh</div>
              <div>Co quan ban hanh Chinh phu</div>
              <div>Trich yeu Sua doi mot so dieu ve thue gia tri gia tang</div>
              <nav>Co quan ban hanh Quoc hoi</nav>
              <a href="https://datafiles.chinhphu.vn/cpp/files/vbpq/2026/05/144-ndcp.signed.pdf">PDF</a>
            </body></html>
            """,
            content_type="text/html",
            headers={},
            fetched_at="2026-05-14T00:00:00+00:00",
            checksum="abc",
            crawl_run_id="run-1",
        )
        record = parse_vanban_chinhphu_detail_record(raw, source)
        self.assertEqual(record.document_type, "LegalDocument")
        self.assertEqual(record.canonical_number, "144/2026/ND-CP")
        self.assertEqual(record.issue_date, "2026-05-05")
        self.assertEqual(record.issuer, "Chinh phu")
        self.assertEqual(record.metadata["crawl_run_id"], "run-1")
        self.assertEqual(record.metadata["attachments"][0]["url"].endswith(".pdf"), True)

    def test_stable_doc_id_and_article_split(self) -> None:
        first = stable_doc_id("vbpl", "59/2020/QH14", "2020-06-17", "Quoc hoi")
        second = stable_doc_id("vbpl", "59-2020-qh14", "2020-06-17", "Quoc hoi")
        self.assertEqual(first, second)
        articles = split_articles("Dieu 1. Title\nBody\n\nDieu 2. Next\nBody")
        self.assertEqual([item["article_number"] for item in articles], ["1", "2"])

    def test_version_store_detects_idempotency_and_changes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            state = DocumentVersionStore(Path(tmp) / "versions.json")
            record = _record(checksum="aaa")
            first = state.plan_update(record, "silver/a.json")
            self.assertIsNotNone(first)
            state.commit(first)

            duplicate = state.plan_update(record, "silver/a.json")
            self.assertIsNone(duplicate)

            changed = state.plan_update(_record(checksum="bbb"), "silver/b.json")
            self.assertIsNotNone(changed)
            self.assertEqual(changed.action, "replace")
            self.assertEqual(changed.previous_version, 1)

    def test_manifest_loader_deduplicates_exact_update_events(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "kg_update_manifest.jsonl"
            manifest = KGUpdateManifest(
                manifest_id="m1",
                doc_id="doc-1",
                source_id="vanban_chinhphu",
                action="insert",
                checksum="abc",
                current_version=1,
                created_at="2026-05-14T00:00:00+00:00",
                source_url="https://vanban.chinhphu.vn/?docid=1&pageid=27160",
                crawl_run_id="run-1",
                silver_record_path="silver.json",
            )
            duplicate = KGUpdateManifest(**{**manifest.to_dict(), "manifest_id": "m2"})
            changed = KGUpdateManifest(**{**manifest.to_dict(), "manifest_id": "m3", "checksum": "def"})
            path.write_text(
                "\n".join(json.dumps(item.to_dict()) for item in (manifest, duplicate, changed)),
                encoding="utf-8",
            )
            items = load_manifest(path)
            self.assertEqual([item.manifest_id for item in items], ["m1", "m3"])

    def test_latest_manifest_items_keeps_last_highest_version_per_doc(self) -> None:
        base = KGUpdateManifest(
            manifest_id="m1",
            doc_id="doc-1",
            source_id="vanban_chinhphu",
            action="insert",
            checksum="abc",
            current_version=1,
            created_at="2026-05-14T00:00:00+00:00",
            source_url="https://vanban.chinhphu.vn/?docid=1&pageid=27160",
            crawl_run_id="run-1",
            silver_record_path="silver-1.json",
        )
        same_doc_later = KGUpdateManifest(
            **{**base.to_dict(), "manifest_id": "m2", "checksum": "def", "silver_record_path": "silver-2.json"}
        )
        other_doc = KGUpdateManifest(**{**base.to_dict(), "manifest_id": "m3", "doc_id": "doc-2"})

        latest = latest_manifest_items([base, same_doc_later, other_doc])

        self.assertEqual([item.manifest_id for item in latest], ["m2", "m3"])

    def test_lightrag_insert_kwargs_match_runtime_signature_without_metadata(self) -> None:
        def insert(input, ids=None, file_paths=None, track_id=None):  # noqa: ANN001
            return None

        record = _record(checksum="abc", crawl_run_id="run-1")
        kwargs = _insert_kwargs(insert, "body", record)

        self.assertEqual(kwargs["input"], ["body"])
        self.assertEqual(kwargs["ids"], [record.doc_id])
        self.assertEqual(kwargs["file_paths"], [record.source_url])
        self.assertEqual(kwargs["track_id"], "run-1")
        self.assertNotIn("metadatas", kwargs)

    def test_document_text_with_provenance_keeps_source_citation(self) -> None:
        record = _record(checksum="abc", crawl_run_id="run-1")
        text = _document_text_with_provenance(record)

        self.assertIn("PROVENANCE", text)
        self.assertIn(f"source_url: {record.source_url}", text)
        self.assertIn(f"doc_id: {record.doc_id}", text)
        self.assertTrue(text.endswith(record.text))

    def test_rerank_settings_are_env_driven_and_safe_by_default(self) -> None:
        with patch.dict(os.environ, {"LIGHTRAG_RERANK_ENABLED": "false"}, clear=False):
            settings = get_rerank_settings()
            self.assertFalse(settings.enabled)
            self.assertIsNone(build_rerank_model_func(settings))

    def test_local_reranker_returns_lightrag_index_format(self) -> None:
        class FakeReranker:
            def predict(self, pairs, batch_size=16, show_progress_bar=False):  # noqa: ANN001
                self.pairs = pairs
                self.batch_size = batch_size
                return [-2.0, 3.0, 0.0]

        settings = RerankSettings(
            enabled=True,
            model_name="fake",
            batch_size=4,
            max_chars_per_doc=12,
            normalize_scores=True,
        )
        fake = FakeReranker()
        import core.rerank_client as rerank_client

        previous = rerank_client._reranker
        rerank_client._reranker = fake
        try:
            func = build_rerank_model_func(settings)
            results = __import__("asyncio").run(
                func("luong", ["abc def", "x" * 50, "bao hiem"], top_n=2)
            )
        finally:
            rerank_client._reranker = previous

        self.assertEqual([item["index"] for item in results], [1, 2])
        self.assertGreater(results[0]["relevance_score"], results[1]["relevance_score"])
        self.assertEqual(fake.batch_size, 4)
        self.assertLessEqual(len(fake.pairs[1][1]), 12)

    def test_local_reranker_writes_score_distribution_trace_when_enabled(self) -> None:
        class FakeReranker:
            def predict(self, pairs, batch_size=16, show_progress_bar=False):  # noqa: ANN001
                return [0.1, 0.7, 0.3]

        settings = RerankSettings(
            enabled=True,
            model_name="fake",
            batch_size=4,
            max_chars_per_doc=100,
            normalize_scores=False,
        )
        fake = FakeReranker()
        import core.rerank_client as rerank_client

        previous = rerank_client._reranker
        rerank_client._reranker = fake
        try:
            with tempfile.TemporaryDirectory() as tmp:
                path = Path(tmp) / "trace.jsonl"
                with patch.dict(
                    os.environ,
                    {
                        "LIGHTRAG_BENCHMARK_TRACE_ENABLED": "true",
                        "LIGHTRAG_BENCHMARK_TRACE_PATH": str(path),
                        "LIGHTRAG_TRACE_PREVIEW_CHARS": "20",
                    },
                    clear=False,
                ):
                    func = build_rerank_model_func(settings)
                    __import__("asyncio").run(func("bao truoc", ["a", "b", "c"], top_n=2))
                event = json.loads(path.read_text(encoding="utf-8").strip())
        finally:
            rerank_client._reranker = previous

        self.assertEqual(event["event_type"], "rerank")
        self.assertEqual(event["candidate_count"], 3)
        self.assertEqual(event["returned_count"], 2)
        self.assertEqual(event["score_summary"]["count"], 3)
        self.assertEqual([item["index"] for item in event["top_results"]], [1, 2])

    def test_rerank_score_helpers_are_bounded(self) -> None:
        self.assertEqual(_truncate_for_rerank("  a   b  ", 20), "a b")
        self.assertLess(_sigmoid(-2.0), 0.5)
        self.assertGreater(_sigmoid(2.0), 0.5)

    def test_benchmark_trace_is_opt_in_and_jsonl(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "trace.jsonl"
            with patch.dict(
                os.environ,
                {
                    "LIGHTRAG_BENCHMARK_TRACE_ENABLED": "false",
                    "LIGHTRAG_BENCHMARK_TRACE_PATH": str(path),
                },
                clear=False,
            ):
                write_trace_event("query", {"query_hash": stable_text_hash("abc")})
                self.assertFalse(path.exists())

            with patch.dict(
                os.environ,
                {
                    "LIGHTRAG_BENCHMARK_TRACE_ENABLED": "true",
                    "LIGHTRAG_BENCHMARK_TRACE_PATH": str(path),
                    "LIGHTRAG_TRACE_PREVIEW_CHARS": "5",
                },
                clear=False,
            ):
                write_trace_event(
                    "query",
                    {
                        "query_hash": stable_text_hash("abc"),
                        "query_preview": text_preview("  abc   def  "),
                    },
                )

            event = json.loads(path.read_text(encoding="utf-8").strip())
            self.assertEqual(event["event_type"], "query")
            self.assertEqual(event["query_preview"], "abc d")
            self.assertEqual(len(event["query_hash"]), 16)

    def test_score_summary_reports_distribution(self) -> None:
        summary = score_summary([0.1, 0.4, 0.2, 0.9])
        self.assertEqual(summary["count"], 4)
        self.assertEqual(summary["min"], 0.1)
        self.assertEqual(summary["max"], 0.9)
        self.assertAlmostEqual(float(summary["p50"]), 0.3)
        self.assertAlmostEqual(float(summary["p90"]), 0.75)

    def test_query_hybrid_uses_env_knobs_and_writes_trace(self) -> None:
        class FakeRag:
            def __init__(self) -> None:
                self.param = None

            async def aquery(self, query, param):  # noqa: ANN001
                self.query = query
                self.param = param
                return "legal context"

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "query_trace.jsonl"
            with patch.dict(
                os.environ,
                {
                    "LIGHTRAG_QUERY_TOP_K": "42",
                    "LIGHTRAG_CHUNK_TOP_K": "21",
                    "LIGHTRAG_RERANK_ENABLED": "false",
                    "LIGHTRAG_BENCHMARK_TRACE_ENABLED": "true",
                    "LIGHTRAG_BENCHMARK_TRACE_PATH": str(path),
                },
                clear=False,
            ):
                rag = FakeRag()
                result = __import__("asyncio").run(query_hybrid(rag, "bao truoc", top_k=10))

            event = json.loads(path.read_text(encoding="utf-8").strip())

        self.assertEqual(result, "legal context")
        self.assertEqual(rag.param.top_k, 42)
        self.assertEqual(rag.param.chunk_top_k, 21)
        self.assertFalse(rag.param.enable_rerank)
        self.assertEqual(event["event_type"], "query")
        self.assertEqual(event["top_k"], 42)
        self.assertEqual(event["chunk_top_k"], 21)
        self.assertFalse(event["enable_rerank"])

    def test_benchmark_runner_builds_isolated_variant_command(self) -> None:
        variant = BenchmarkVariant("off", {"LIGHTRAG_RERANK_ENABLED": "false"})
        command, env, paths = build_variant_command(
            variant,
            groundtruth=Path("gt.json"),
            run_dir=Path("reports/benchmarks/run-1"),
            contract=Path("contract.docx"),
        )

        self.assertIn("e2e_eval.py", command[1])
        self.assertIn("--contract", command)
        self.assertEqual(env["LIGHTRAG_RERANK_ENABLED"], "false")
        self.assertEqual(env["LIGHTRAG_BENCHMARK_TRACE_ENABLED"], "true")
        self.assertTrue(str(paths["trace"]).endswith("off\\retrieval_trace.jsonl") or str(paths["trace"]).endswith("off/retrieval_trace.jsonl"))

    def test_benchmark_runner_parses_eval_state_and_trace_summaries(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            base = Path(tmp)
            eval_path = base / "eval.md"
            eval_path.write_text(
                "\n".join(
                    [
                        "- Predicted violations: 2",
                        "- Groundtruth vulnerabilities: 5",
                        "- Precision (heuristic): 0.500",
                        "- Recall (heuristic): 0.200",
                        "- F1 (heuristic): 0.286",
                    ]
                ),
                encoding="utf-8",
            )
            state_path = base / "state.json"
            state_path.write_text(
                json.dumps(
                    {
                        "context_quality": "good",
                        "context_quality_score": 0.75,
                        "context_retry_count": 0,
                        "error_type": "reasoning",
                        "confidence": 0.6,
                        "audit_findings": [{"id": 1}, {"id": 2}],
                        "retrieved_clause_indices": [1, 3],
                        "critic_feedback": {
                            "findings_pruned": 4,
                            "admissibility_reason": "#0:reference_not_supported",
                        },
                    }
                ),
                encoding="utf-8",
            )
            trace_path = base / "trace.jsonl"
            trace_path.write_text(
                "\n".join(
                    [
                        json.dumps({"event_type": "query"}),
                        json.dumps({"event_type": "rerank", "candidate_count": 12, "score_summary": {"p50": 0.4, "max": 0.8}}),
                        json.dumps({"event_type": "rerank", "candidate_count": 20, "score_summary": {"p50": 0.6, "max": 0.9}}),
                    ]
                ),
                encoding="utf-8",
            )

            metrics = parse_eval_metrics(eval_path)
            state = parse_state_summary(state_path)
            trace = parse_trace_summary(trace_path)

        self.assertEqual(metrics["predicted_violations"], 2)
        self.assertEqual(metrics["recall"], 0.2)
        self.assertEqual(state["critic_findings_pruned"], 4)
        self.assertEqual(trace["rerank_events"], 2)
        self.assertEqual(trace["rerank_candidate_max"], 20)
        self.assertAlmostEqual(trace["rerank_score_p50_avg"], 0.5)

    def test_benchmark_summary_markdown_compares_variants(self) -> None:
        markdown = build_summary_markdown(
            [
                {
                    "variant": "baseline_no_rerank",
                    "returncode": 0,
                    "metrics": {"precision": 1.0, "recall": 0.2, "f1": 0.333},
                    "state_summary": {"critic_findings_pruned": 7},
                    "trace_summary": {"rerank_events": 0},
                }
            ]
        )

        self.assertIn("baseline_no_rerank", markdown)
        self.assertIn("0.200", markdown)
        self.assertIn("Rerank events", markdown)

    def test_benchmark_diagnosis_flags_critic_prune_loss(self) -> None:
        diagnosis = diagnose_result(
            {
                "variant": "rerank_default",
                "returncode": 0,
                "env": {"LIGHTRAG_RERANK_ENABLED": "true"},
                "metrics": {
                    "predicted_violations": 1,
                    "groundtruth_vulnerabilities": 5,
                    "recall": 0.2,
                },
                "state_summary": {
                    "context_quality": "good",
                    "context_quality_score": 0.75,
                    "critic_findings_pruned": 7,
                    "critic_admissibility_reason": "#0:reference_not_supported",
                },
                "trace_summary": {"rerank_events": 11},
            }
        )

        self.assertEqual(diagnosis["likely_loss_point"], "critic_admissibility_prune")
        self.assertIn("Critic pruned 7 findings", " ".join(diagnosis["findings"]))

    def test_benchmark_diagnosis_compares_rerank_variants(self) -> None:
        diagnosis = build_diagnosis(
            {
                "run_dir": "reports/benchmarks/run",
                "results": [
                    {
                        "variant": "baseline_no_rerank",
                        "returncode": 0,
                        "env": {"LIGHTRAG_RERANK_ENABLED": "false"},
                        "metrics": {"recall": 0.4, "predicted_violations": 2, "groundtruth_vulnerabilities": 5},
                        "state_summary": {"context_quality": "good", "context_quality_score": 0.8},
                        "trace_summary": {"rerank_events": 0},
                    },
                    {
                        "variant": "rerank_default",
                        "returncode": 0,
                        "env": {"LIGHTRAG_RERANK_ENABLED": "true"},
                        "metrics": {"recall": 0.2, "predicted_violations": 1, "groundtruth_vulnerabilities": 5},
                        "state_summary": {"context_quality": "good", "context_quality_score": 0.8},
                        "trace_summary": {"rerank_events": 10},
                    },
                ],
            }
        )
        markdown = build_diagnosis_markdown(diagnosis)

        self.assertEqual(diagnosis["variant_comparison"]["rerank_effect"], "rerank_reduced_recall")
        self.assertIn("Increase first-stage candidate pool", "\n".join(diagnosis["next_actions"]))
        self.assertIn("rerank_reduced_recall", markdown)

    def test_benchmark_runner_has_candidate_pool_sweep_variants_without_threshold(self) -> None:
        self.assertEqual([variant.name for variant in SWEEP_VARIANTS], ["rerank_wide_30_40", "rerank_wide_50_60"])
        self.assertTrue(all(variant.env["LIGHTRAG_MIN_RERANK_SCORE"] == "0.0" for variant in SWEEP_VARIANTS))
        self.assertEqual(SWEEP_VARIANTS[0].env["LIGHTRAG_QUERY_TOP_K"], "30")

    def test_calibration_recommends_candidate_pool_when_rerank_reduces_recall(self) -> None:
        calibration = build_calibration(
            summary={"results": [{"variant": "rerank_default", "metrics": {"recall": 0.2}}]},
            diagnosis={
                "variant_comparison": {"rerank_effect": "rerank_reduced_recall"},
                "diagnoses": [],
            },
        )
        markdown = build_calibration_markdown(calibration)

        self.assertEqual(calibration["decision"], "increase_candidate_pool_before_threshold")
        self.assertEqual(calibration["recommended_env"]["LIGHTRAG_MIN_RERANK_SCORE"], "0.0")
        self.assertEqual(calibration["recommended_env"]["LIGHTRAG_QUERY_TOP_K"], "30")
        self.assertIn("--include-sweep", calibration["next_benchmark_command_flags"])
        self.assertIn("LIGHTRAG_QUERY_TOP_K=30", markdown)

    def test_calibration_prioritizes_critic_debug_over_threshold(self) -> None:
        calibration = build_calibration(
            summary={"results": []},
            diagnosis={
                "variant_comparison": {"rerank_effect": "rerank_neutral_on_recall"},
                "diagnoses": [{"likely_loss_point": "critic_admissibility_prune"}],
            },
        )

        self.assertEqual(calibration["decision"], "debug_critic_before_retrieval_tuning")
        self.assertEqual(calibration["recommended_env"]["LIGHTRAG_MIN_RERANK_SCORE"], "0.0")

    def test_docx_attachment_extraction_and_record_merge_updates_checksum(self) -> None:
        import docx

        buffer = BytesIO()
        document = docx.Document()
        document.add_paragraph("Dieu 1. Noi dung file dinh kem chinh thuc.")
        document.add_paragraph("Dieu 2. Van ban nay duoc extract tu DOCX.")
        document.save(buffer)
        raw = RawArtifact(
            source_item=SourceItem(
                source_id="vanban_chinhphu",
                url="https://datafiles.chinhphu.vn/example.docx",
                title="example.docx",
                metadata={"artifact_kind": "attachment", "parent_doc_id": "doc-1"},
            ),
            content=buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={},
            fetched_at="2026-05-14T00:00:00+00:00",
            checksum=checksum_bytes(buffer.getvalue()),
            crawl_run_id="run-1",
        )
        extraction = extract_attachment_text(raw, min_chars=20)
        self.assertEqual(extraction.status, "extracted")
        self.assertIn("Noi dung file dinh kem", extraction.text)

        record = _record(checksum="html-checksum", crawl_run_id="run-1")
        merged = merge_attachment_extractions(record, [extraction], min_chars=20)
        self.assertNotEqual(merged.checksum, record.checksum)
        self.assertIn("NOI DUNG FILE DINH KEM CHINH THUC", merged.text)
        self.assertEqual(merged.metadata["attachment_extraction"]["extracted_count"], 1)

    def test_blank_pdf_attachment_is_marked_for_ocr_review(self) -> None:
        from pypdf import PdfWriter

        buffer = BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        writer.write(buffer)
        raw = RawArtifact(
            source_item=SourceItem(
                source_id="vanban_chinhphu",
                url="https://datafiles.chinhphu.vn/blank.pdf",
                title="blank.pdf",
                metadata={"artifact_kind": "attachment", "parent_doc_id": "doc-1"},
            ),
            content=buffer.getvalue(),
            content_type="application/pdf",
            headers={},
            fetched_at="2026-05-14T00:00:00+00:00",
            checksum=checksum_bytes(buffer.getvalue()),
            crawl_run_id="run-1",
        )
        extraction = extract_attachment_text(raw, min_chars=20)
        self.assertEqual(extraction.status, "needs_ocr")
        self.assertTrue(extraction.needs_review)

    def test_pdf_attachment_can_fallback_to_tika_ocr(self) -> None:
        from pypdf import PdfWriter

        buffer = BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        writer.write(buffer)
        raw = RawArtifact(
            source_item=SourceItem(
                source_id="vanban_chinhphu",
                url="https://datafiles.chinhphu.vn/scanned.pdf",
                title="scanned.pdf",
                metadata={"artifact_kind": "attachment", "parent_doc_id": "doc-1"},
            ),
            content=buffer.getvalue(),
            content_type="application/pdf",
            headers={},
            fetched_at="2026-05-14T00:00:00+00:00",
            checksum=checksum_bytes(buffer.getvalue()),
            crawl_run_id="run-1",
        )

        class FakeResponse:
            text = "Dieu 1. Noi dung OCR tu Tika cho van ban scan."

            def raise_for_status(self) -> None:
                return None

        with patch("pipeline.attachment_extraction.requests.put", return_value=FakeResponse()) as mocked:
            extraction = extract_attachment_text(raw, min_chars=20, tika_server_url="http://tika:9998")
        self.assertEqual(extraction.status, "extracted")
        self.assertEqual(extraction.method, "tika_ocr")
        self.assertIn("Noi dung OCR", extraction.text)
        self.assertEqual(mocked.call_args.kwargs["headers"]["X-Tika-OCRLanguage"], "vie+eng")

    def test_attachment_fetch_blocks_hosts_outside_whitelist(self) -> None:
        item = SourceItem(
            source_id="vanban_chinhphu",
            url="https://evil.example/file.pdf",
            title="file.pdf",
        )
        with self.assertRaises(ValueError):
            fetch_attachment(item, allowed_hosts={"vanban.chinhphu.vn", "datafiles.chinhphu.vn"})

    def test_local_lakehouse_writes_layers(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            lakehouse = LocalLakehouse(tmp)
            raw = RawArtifact(
                source_item=SourceItem(source_id="vbpl", url="https://vbpl.vn/a", title="A"),
                content=b"content",
                content_type="text/html",
                headers={"x": "y"},
                fetched_at="2026-05-14T00:00:00+00:00",
                checksum=checksum_bytes(b"content"),
                crawl_run_id="run-1",
            )
            bronze_path = lakehouse.write_bronze(raw)
            silver_path = lakehouse.write_silver(_record(checksum=raw.checksum))
            self.assertTrue(bronze_path.exists())
            self.assertIn("run-1", str(bronze_path))
            self.assertEqual(json.loads(silver_path.read_text(encoding="utf-8"))["checksum"], raw.checksum)

    def test_iceberg_rows_are_flat_and_debuggable(self) -> None:
        raw = RawArtifact(
            source_item=SourceItem(
                source_id="vbpl",
                url="https://vbpl.vn/a",
                title="A",
                metadata={"feed_id": "fixture"},
            ),
            content=b"content",
            content_type="text/html",
            headers={"x": "y"},
            fetched_at="2026-05-14T00:00:00+00:00",
            checksum=checksum_bytes(b"content"),
            crawl_run_id="run-1",
        )
        record = _record(checksum=raw.checksum, crawl_run_id="run-1")
        manifest = KGUpdateManifest(
            manifest_id="m1",
            doc_id=record.doc_id,
            source_id=record.source_id,
            action="insert",
            checksum=record.checksum,
            current_version=1,
            created_at="2026-05-14T00:00:00+00:00",
            source_url=record.source_url,
            silver_record_path="silver.json",
            crawl_run_id="run-1",
        )
        self.assertEqual(bronze_row(raw)["content"], b"content")
        self.assertEqual(bronze_row(raw)["crawl_run_id"], "run-1")
        self.assertEqual(silver_row(record)["crawl_run_id"], "run-1")
        self.assertEqual(gold_row(manifest)["crawl_run_id"], "run-1")
        self.assertEqual(json.loads(silver_row(record)["metadata_json"]), record.metadata)
        self.assertEqual(gold_row(manifest)["current_version"], 1)

    def test_iceberg_uri_fallback_preserves_credentials(self) -> None:
        uri = "postgresql+psycopg2://rag:rag_secure_pwd@postgres:5432/viet_contract"
        replaced = _replace_netloc(uri, "127.0.0.1", 5433)
        self.assertEqual(
            replaced,
            "postgresql+psycopg2://rag:rag_secure_pwd@127.0.0.1:5433/viet_contract",
        )

    def test_uri_redaction_hides_password(self) -> None:
        uri = "postgresql+psycopg2://rag:rag_secure_pwd@127.0.0.1:5433/viet_contract"
        self.assertEqual(
            redact_uri(uri),
            "postgresql+psycopg2://rag:***@127.0.0.1:5433/viet_contract",
        )

    def test_source_item_external_id_is_used_for_record_doc_id(self) -> None:
        from pipeline.connectors import LegalSourceConnector

        class FixtureConnector(LegalSourceConnector):
            def discover(self, since=None, limit=None):
                return []

        source = SourceConfig(
            source_id="congbao",
            name="Cong bao",
            tier=0,
            role="canonical_update_feed",
            base_url="https://congbao.chinhphu.vn",
            domain_whitelist=["congbao.chinhphu.vn"],
            crawl_methods=["rss"],
            rate_limit_per_minute=30,
            robots_policy="obey",
            priority=10,
            license_note="attribute",
        )
        raw = RawArtifact(
            source_item=SourceItem(
                source_id="congbao",
                url="https://congbao.chinhphu.vn/doc",
                title="A",
                external_id="stable-guid",
            ),
            content=b"<html>A</html>",
            content_type="text/html",
            headers={},
            fetched_at="2026-05-14T00:00:00+00:00",
            checksum="abc",
            crawl_run_id="run-1",
        )
        record = FixtureConnector(source).parse(raw)
        self.assertEqual(record.canonical_number, "stable-guid")

    def test_cong_bao_issue_feed_is_not_legal_document(self) -> None:
        from pipeline.connectors import LegalSourceConnector

        class FixtureConnector(LegalSourceConnector):
            def discover(self, since=None, limit=None):
                return []

        source = SourceConfig(
            source_id="congbao",
            name="Cong bao",
            tier=0,
            role="canonical_update_feed",
            base_url="https://congbao.chinhphu.vn",
            domain_whitelist=["congbao.chinhphu.vn"],
            crawl_methods=["rss"],
            rate_limit_per_minute=30,
            robots_policy="obey",
            priority=10,
            license_note="attribute",
        )
        raw = RawArtifact(
            source_item=SourceItem(
                source_id="congbao",
                url="https://congbao.chinhphu.vn/cong-bao/cong-bao-so-1.htm",
                title="Cong bao so 1",
                metadata={"feed_id": "cong_bao_moi_dang"},
            ),
            content=b"<html>Cong bao so 1</html>",
            content_type="text/html",
            headers={},
            fetched_at="2026-05-14T00:00:00+00:00",
            checksum="abc",
            crawl_run_id="run-1",
        )
        record = FixtureConnector(source).parse(raw)
        self.assertEqual(record.document_type, "OfficialGazetteIssue")


def _record(checksum: str, crawl_run_id: str | None = None) -> LegalDocumentRecord:
    return LegalDocumentRecord(
        doc_id="vbpl:doc",
        source_id="vbpl",
        canonical_number="59/2020/QH14",
        issue_date="2020-06-17",
        issuer="Quoc hoi",
        title="Luật Doanh nghiệp",
        text="Dieu 1. Text",
        source_url="https://vbpl.vn/doc",
        checksum=checksum,
        fetched_at="2026-05-14T00:00:00+00:00",
        crawl_run_id=crawl_run_id,
    )


if __name__ == "__main__":
    unittest.main()
