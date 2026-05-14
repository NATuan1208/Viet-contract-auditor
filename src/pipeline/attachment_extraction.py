"""Official attachment download and text extraction helpers."""

from __future__ import annotations

import io
import json
from dataclasses import asdict, dataclass, replace
from pathlib import PurePosixPath
from typing import Any
from urllib.parse import urlparse

import requests

from .models import LegalDocumentRecord, RawArtifact, SourceItem, utc_now_iso
from .normalization import checksum_bytes, clean_text, split_articles

DEFAULT_ATTACHMENT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (compatible; VietContractAuditor/1.0; "
        "+https://github.com/local/viet-contract-auditor)"
    )
}


@dataclass(frozen=True)
class AttachmentExtractionResult:
    source_url: str
    title: str
    checksum: str
    method: str
    status: str
    text: str
    char_count: int
    needs_review: bool
    error: str = ""

    def to_dict(self) -> dict[str, Any]:
        payload = asdict(self)
        payload["text_preview"] = self.text[:500]
        payload.pop("text", None)
        return payload


def attachment_item(
    parent: LegalDocumentRecord,
    attachment: dict[str, str],
    index: int,
) -> SourceItem:
    url = attachment.get("url", "")
    title = attachment.get("title") or _filename_from_url(url) or f"attachment-{index}"
    return SourceItem(
        source_id=parent.source_id,
        url=url,
        title=title,
        external_id=f"{parent.doc_id}:attachment:{index}",
        metadata={
            "artifact_kind": "attachment",
            "parent_doc_id": parent.doc_id,
            "parent_source_url": parent.source_url,
            "attachment_index": index,
            "attachment_title": title,
        },
    )


def fetch_attachment(
    item: SourceItem,
    crawl_run_id: str | None = None,
    timeout_seconds: int = 30,
    allowed_hosts: set[str] | None = None,
) -> RawArtifact:
    host = urlparse(item.url).hostname
    if allowed_hosts is not None and host not in allowed_hosts:
        raise ValueError(f"{item.url} is outside attachment host whitelist")
    response = requests.get(item.url, timeout=timeout_seconds, headers=DEFAULT_ATTACHMENT_HEADERS)
    response.raise_for_status()
    content = response.content
    return RawArtifact(
        source_item=item,
        content=content,
        content_type=response.headers.get("content-type", "application/octet-stream"),
        headers=dict(response.headers),
        fetched_at=utc_now_iso(),
        checksum=checksum_bytes(content),
        crawl_run_id=crawl_run_id,
    )


def extract_attachment_text(
    raw: RawArtifact,
    min_chars: int = 200,
    tika_server_url: str | None = None,
) -> AttachmentExtractionResult:
    suffix = _suffix(raw.source_item.url)
    try:
        if suffix == ".pdf" or "pdf" in raw.content_type.lower():
            text = _extract_pdf_text(raw.content)
            method = "pypdf_text_layer"
            if len(text) < min_chars and tika_server_url:
                text = _extract_with_tika(raw, tika_server_url, ocr=True)
                return _result(raw, method="tika_ocr", status="extracted", text=text, min_chars=min_chars)
            status = "extracted" if len(text) >= min_chars else "needs_ocr"
            return _result(raw, method=method, status=status, text=text, min_chars=min_chars)
        if suffix == ".docx":
            text = _extract_docx_text(raw.content)
            return _result(raw, method="python_docx", status="extracted", text=text, min_chars=min_chars)
        if suffix in {".doc", ".rtf"} and tika_server_url:
            text = _extract_with_tika(raw, tika_server_url, ocr=False)
            return _result(raw, method="tika_legacy_office", status="extracted", text=text, min_chars=min_chars)
        if suffix in {".doc", ".rtf"}:
            return _result(raw, method="unsupported_legacy_office", status="unsupported", text="", min_chars=min_chars)
        return _result(raw, method="unsupported_format", status="unsupported", text="", min_chars=min_chars)
    except Exception as exc:
        return AttachmentExtractionResult(
            source_url=raw.source_item.url,
            title=raw.source_item.title,
            checksum=raw.checksum,
            method="error",
            status="error",
            text="",
            char_count=0,
            needs_review=True,
            error=str(exc),
        )


def merge_attachment_extractions(
    record: LegalDocumentRecord,
    extractions: list[AttachmentExtractionResult],
    min_chars: int = 200,
) -> LegalDocumentRecord:
    usable = [item for item in extractions if item.status == "extracted" and item.char_count >= min_chars]
    extraction_summary = {
        "attempted_count": len(extractions),
        "extracted_count": len(usable),
        "needs_review_count": sum(1 for item in extractions if item.needs_review),
        "items": [item.to_dict() for item in extractions],
    }
    metadata = {
        **record.metadata,
        "attachment_extraction": extraction_summary,
    }
    if not usable:
        return replace(record, metadata=metadata)

    attachment_sections = []
    for index, item in enumerate(usable, start=1):
        attachment_sections.append(
            "\n".join(
                (
                    f"NOI DUNG FILE DINH KEM CHINH THUC {index}: {item.title}",
                    f"Nguon file: {item.source_url}",
                    item.text,
                )
            )
        )
    merged_text = record.text + "\n\n" + "\n\n".join(attachment_sections)
    checksum_payload = {
        "html_checksum": record.checksum,
        "attachments": [
            {
                "source_url": item.source_url,
                "checksum": item.checksum,
                "method": item.method,
                "char_count": item.char_count,
            }
            for item in usable
        ],
    }
    return replace(
        record,
        text=merged_text,
        checksum=checksum_bytes(json.dumps(checksum_payload, sort_keys=True).encode("utf-8")),
        articles=split_articles(merged_text),
        metadata=metadata,
    )


def _extract_pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return clean_text("\n\n".join(pages))


def _extract_docx_text(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
    return clean_text("\n".join(paragraphs + table_lines))


def _extract_with_tika(raw: RawArtifact, tika_server_url: str, ocr: bool) -> str:
    headers = {
        "Accept": "text/plain",
        "Content-Type": raw.content_type or "application/octet-stream",
        "X-Tika-OCRLanguage": "vie+eng",
    }
    if ocr:
        headers["X-Tika-PDFOcrStrategy"] = "ocr_only"
    response = requests.put(
        f"{tika_server_url.rstrip('/')}/tika",
        data=raw.content,
        headers=headers,
        timeout=180,
    )
    response.raise_for_status()
    return clean_text(response.text)


def _result(
    raw: RawArtifact,
    method: str,
    status: str,
    text: str,
    min_chars: int,
) -> AttachmentExtractionResult:
    text = clean_text(text)
    char_count = len(text)
    needs_review = status != "extracted" or char_count < min_chars
    if status == "extracted" and char_count < min_chars:
        status = "too_short"
    return AttachmentExtractionResult(
        source_url=raw.source_item.url,
        title=raw.source_item.title,
        checksum=raw.checksum,
        method=method,
        status=status,
        text=text,
        char_count=char_count,
        needs_review=needs_review,
    )


def _suffix(url: str) -> str:
    return PurePosixPath(urlparse(url).path).suffix.lower()


def _filename_from_url(url: str) -> str:
    return PurePosixPath(urlparse(url).path).name
